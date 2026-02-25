"""
Transcriber Pro — Local Review Dashboard
Lightweight local server for reviewing cloud-transcribed results.
No GPU needed. Loads .json state files and pairs them with local audio.

Also supports S3 upload for sending files to RunPod.
"""

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import os
import sys
import io
import json
import re
import threading
import time
import uuid
from pathlib import Path
import requests as http_requests
import paramiko
from scp import SCPClient
import tarfile
import shutil

import boto3
from botocore.config import Config
from docx import Document
from dotenv import load_dotenv

# Load .env credentials
load_dotenv()

# Fix UTF-8 on Windows
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Directories ───
UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

# ─── S3 Config ───
S3_BUCKET = "ez2d4o9xmt"
S3_ENDPOINT = "https://s3api-us-wa-1.runpod.io"
S3_REGION = "us-wa-1"

s3 = boto3.client(
    "s3",
    endpoint_url=S3_ENDPOINT,
    region_name=S3_REGION,
    aws_access_key_id=os.getenv("RUNPOD_ACCESS_KEY"),
    aws_secret_access_key=os.getenv("RUNPOD_SECRET_KEY"),
    config=Config(signature_version="s3v4"),
)

# ─── RunPod API & SSH Config ───
RUNPOD_API_KEY = os.getenv("RUNPOD_API_KEY")
RUNPOD_POD_ID = os.getenv("RUNPOD_POD_ID", "")
RUNPOD_ENDPOINT_ID = os.getenv("RUNPOD_ENDPOINT_ID", "") # New Serverless Endpoint
RUNPOD_GQL = "https://api.runpod.io/graphql"
HF_TOKEN = os.getenv("HF_TOKEN", "") # HuggingFace Token for Diarization

# SSH Config (User-provided via UI)
pod_config = {
    "ip": os.getenv("POD_IP", ""),
    "ssh_port": int(os.getenv("POD_SSH_PORT", "22")),
    "key_path": os.getenv("POD_KEY_PATH", "runpod")
}

# ─── State ───
transcriptions = {}

def download_results_from_s3():
    """Check S3 for any finished results (.json, .md, .docx) and pull them to local uploads."""
    try:
        response = s3.list_objects_v2(Bucket=S3_BUCKET)
        if 'Contents' not in response:
            return

        result_exts = {".json", ".md", ".docx"}
        found_new = False
        for obj in response['Contents']:
            s3_key = obj['Key']
            # Strip the 'transcriber/uploads/' prefix if present
            base_name = s3_key.split('/')[-1] if '/' in s3_key else s3_key
            ext = Path(base_name).suffix.lower()

            if ext in result_exts:
                local_path = UPLOAD_DIR / base_name
                if not local_path.exists():
                    if not found_new:
                        print("☁️ New results found on cloud!")
                        found_new = True
                    print(f"  📥 Downloading: {base_name}")
                    s3.download_file(S3_BUCKET, s3_key, str(local_path))

                    if ext == ".json":
                        try:
                            with open(local_path, "r", encoding="utf-8") as f:
                                data = json.load(f)
                                task_id = data.get("filename")
                                if task_id:
                                    transcriptions[task_id] = data
                        except:
                            pass
    except Exception as e:
        print(f"⚠️ Cloud sync check failed: {e}")

def load_existing_tasks():
    """Load previously completed transcriptions from JSON files on disk."""
    print("📂 Scanning local uploads...")
    for json_file in UPLOAD_DIR.glob("*.json"):
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                task_id = data.get("filename")
                if task_id:
                    transcriptions[task_id] = data
                    print(f"  ✅ Loaded: {task_id}")
        except Exception as e:
            print(f"  ❌ Failed to load {json_file.name}: {e}")

# Initial load from disk
load_existing_tasks()

# Start a background thread to check cloud every 30 seconds
def cloud_watchdog():
    # Initial sync on first run (non-blocking)
    download_results_from_s3()
    while True:
        time.sleep(30)
        download_results_from_s3()

threading.Thread(target=cloud_watchdog, daemon=True).start()


# ─── Helpers ───

def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02}:{minutes:02}:{secs:02}"
    return f"{minutes:02}:{secs:02}"


def clean_hallucinations(text: str) -> str:
    """Remove common Russian Whisper hallucinations."""
    hallucination_patterns = [
        r'\bРедактор субтитров\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bКорректор\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bСубтитры\s*:\s*[^\.]+',
        r'\bПеревод\s*:\s*[^\.]+',
        r'\bОзвучка\s*:\s*[^\.]+',
        r'\bРедактор субтитров\b',
        r'\bКорректор\b',
        r'\b(Все права защищены|Продолжение следует|Ставьте лайки|Подписывайтесь на канал)\b',
    ]
    cleaned = text
    for pattern in hallucination_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def generate_docx(task_id):
    """Generate a .docx file from the transcription segments."""
    if task_id not in transcriptions:
        return None
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]
    docx_file_path = file_path.with_suffix(".docx")

    doc = Document()
    doc.add_heading(f"Transcription: {task['filename']}", 0)
    for seg in task["result"]:
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{seg['timestamp']}] {seg['speaker']}: ")
        ts_run.bold = True
        p.add_run(seg['text'])
    doc.save(docx_file_path)
    return docx_file_path.name


def regenerate_files(task_id):
    """Re-save .md, .docx, and .json after speaker edits."""
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]

    # MD
    md_file_path = file_path.with_suffix(".md")
    md_content = f"# Transcription: {task['filename']}\n\n"
    for seg in task["result"]:
        md_content += f"**[{seg['timestamp']}] {seg['speaker']}:** {seg['text']}\n\n"
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    # DOCX
    generate_docx(task_id)

    # JSON state
    state_file = file_path.with_suffix(".json")
    with open(state_file, "w", encoding="utf-8") as f:
        json.dump(task, f, indent=2, ensure_ascii=False)


# ─── S3 Upload (Background Thread) ───

def upload_to_s3(file_path: Path, task_id: str):
    """Upload audio file to RunPod S3 bucket in background."""
    try:
        transcriptions[task_id]["status"] = "uploading"
        transcriptions[task_id]["progress"] = 5

        file_size = file_path.stat().st_size
        uploaded = 0

        def progress_callback(bytes_transferred):
            nonlocal uploaded
            uploaded += bytes_transferred
            pct = min(int((uploaded / file_size) * 90), 90)
            transcriptions[task_id]["progress"] = pct

        from boto3.s3.transfer import TransferConfig
        
        # Use a safe ASCII key for S3 to prevent URL encoding mismatch with Signature v4
        safe_key = f"{uuid.uuid4().hex}_{int(time.time())}{file_path.suffix}"

        s3.upload_file(
            str(file_path),
            S3_BUCKET,
            f"transcriber/uploads/{safe_key}",
            Callback=progress_callback,
            Config=TransferConfig(multipart_threshold=2 * 1024 * 1024 * 1024)
        )

        transcriptions[task_id]["status"] = "uploaded"
        transcriptions[task_id]["progress"] = 100
        transcriptions[task_id]["s3_key"] = safe_key
        print(f"☁️ Uploaded {file_path.name} to S3 as {safe_key}")

    except Exception as e:
        transcriptions[task_id]["status"] = "error"
        transcriptions[task_id]["error"] = f"S3 upload failed: {e}"
        print(f"❌ S3 upload failed: {e}")


@app.post("/diarize-cloud/{task_id}")
async def diarize_cloud(task_id: str, min_speakers: int = 2, max_speakers: int = 6):
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    task = transcriptions[task_id]
    if not RUNPOD_ENDPOINT_ID:
        return {"error": "RUNPOD_ENDPOINT_ID not set in .env"}

    def poll_job(job_id, task_id):
        headers = {
            "Authorization": f"Bearer {RUNPOD_API_KEY}",
            "Content-Type": "application/json"
        }
        status_url = f"https://api.runpod.ai/v2/{RUNPOD_ENDPOINT_ID}/status/{job_id}"
        
        while True:
            try:
                resp = http_requests.get(status_url, headers=headers)
                data = resp.json()
                status = data.get("status")
                
                if status == "COMPLETED":
                    output = data["output"]
                    timeline = output.get("timeline", [])
                    
                    transcriptions[task_id]["timeline"] = timeline
                    transcriptions[task_id]["status"] = "diarization_complete"
                    transcriptions[task_id]["progress"] = 100
                    
                    # Cache the diarization back to JSON
                    json_path = UPLOAD_DIR / Path(task_id).with_suffix(".json")
                    with open(json_path, "w", encoding="utf-8") as f:
                        json.dump(transcriptions[task_id], f, indent=2, ensure_ascii=False)
                        
                    print(f"✅ Serverless Diarization Done: {task_id}")
                    break
                elif status in ["FAILED", "CANCELLED"]:
                    error_msg = data.get("error", "Job failed")
                    transcriptions[task_id]["status"] = "error"
                    transcriptions[task_id]["error"] = error_msg
                    print(f"❌ Serverless Job Failed ({job_id}): {error_msg}")
                    break
                
                if status == "IN_PROGRESS":
                    transcriptions[task_id]["status"] = "diarizing"
                    transcriptions[task_id]["progress"] = 50
                elif status == "IN_QUEUE":
                    transcriptions[task_id]["status"] = "diarizing"
                    transcriptions[task_id]["progress"] = 20

                time.sleep(5)
            except Exception as e:
                print(f"⚠️ Error polling job {job_id}: {e}")
                time.sleep(10)

    try:
        safe_key = task.get("s3_key", task_id)
        s3_key = f"transcriber/uploads/{safe_key}"
        url = f"https://api.runpod.ai/v2/{RUNPOD_ENDPOINT_ID}/run"
        headers = {
            "Authorization": f"Bearer {RUNPOD_API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "input": {
                "action": "diarize",
                "audio": s3_key,
                "s3_creds": {
                    "endpoint": S3_ENDPOINT,
                    "region": S3_REGION,
                    "access_key": os.getenv("RUNPOD_ACCESS_KEY"),
                    "secret_key": os.getenv("RUNPOD_SECRET_KEY"),
                    "bucket": S3_BUCKET
                },
                "min_speakers": min_speakers,
                "max_speakers": max_speakers,
                "hf_token": HF_TOKEN
            }
        }
        
        resp = http_requests.post(url, headers=headers, json=payload)
        resp_data = resp.json()
        job_id = resp_data.get("id")
        
        if job_id:
            task["status"] = "diarizing"
            task["progress"] = 10
            task["job_id"] = job_id
            
            threading.Thread(target=poll_job, args=(job_id, task_id), daemon=True).start()
            print(f"🚀 Serverless Diarization Job Started: {job_id} for {task_id}")
            return {"status": "started", "job_id": job_id}
        else:
            return {"status": "error", "error": f"Failed to start job: {resp_data}"}

    except Exception as e:
        task["status"] = "error"
        task["error"] = str(e)
        return {"status": "error", "error": str(e)}

@app.post("/transcribe-cloud/{task_id}")
async def transcribe_cloud(task_id: str):
    """Trigger RunPod Serverless transcription using background polling."""
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    task = transcriptions[task_id]
    if not RUNPOD_ENDPOINT_ID:
        return {"error": "RUNPOD_ENDPOINT_ID not set in .env"}

    def poll_job(job_id, task_id):
        headers = {
            "Authorization": f"Bearer {RUNPOD_API_KEY}",
            "Content-Type": "application/json"
        }
        status_url = f"https://api.runpod.ai/v2/{RUNPOD_ENDPOINT_ID}/status/{job_id}"
        
        while True:
            try:
                resp = http_requests.get(status_url, headers=headers)
                data = resp.json()
                status = data.get("status")
                
                if status == "COMPLETED":
                    output = data["output"]
                    formatted_segments = []
                    for seg in output.get("result", []):
                        formatted_segments.append({
                            "start": seg["start"],
                            "end": seg.get("end", seg["start"] + 2),
                            "timestamp": format_timestamp(seg["start"]),
                            "speaker": seg.get("speaker", "Unknown"),
                            "text": seg["text"]
                        })
                    
                    transcriptions[task_id]["result"] = formatted_segments
                    transcriptions[task_id]["status"] = "completed"
                    transcriptions[task_id]["progress"] = 100
                    regenerate_files(task_id)
                    print(f"✅ Serverless Transcription Done: {task_id}")
                    break
                elif status in ["FAILED", "CANCELLED"]:
                    error_msg = data.get("error", "Job failed")
                    transcriptions[task_id]["status"] = "error"
                    transcriptions[task_id]["error"] = error_msg
                    print(f"❌ Serverless Job Failed ({job_id}): {error_msg}")
                    break
                
                if status == "IN_PROGRESS":
                    transcriptions[task_id]["status"] = "transcribing"
                    transcriptions[task_id]["progress"] = 50
                elif status == "IN_QUEUE":
                    transcriptions[task_id]["status"] = "transcribing"
                    transcriptions[task_id]["progress"] = 20

                time.sleep(5)
            except Exception as e:
                print(f"⚠️ Error polling job {job_id}: {e}")
                time.sleep(10)

    try:
        safe_key = task.get("s3_key", task_id)
        s3_key = f"transcriber/uploads/{safe_key}"
        
        url = f"https://api.runpod.ai/v2/{RUNPOD_ENDPOINT_ID}/run"
        headers = {
            "Authorization": f"Bearer {RUNPOD_API_KEY}",
            "Content-Type": "application/json"
        }
        
        timeline = task.get("timeline", [])
        
        payload = {
            "input": {
                "action": "transcribe",
                "audio": s3_key,
                "s3_creds": {
                    "endpoint": S3_ENDPOINT,
                    "region": S3_REGION,
                    "access_key": os.getenv("RUNPOD_ACCESS_KEY"),
                    "secret_key": os.getenv("RUNPOD_SECRET_KEY"),
                    "bucket": S3_BUCKET
                },
                "timeline": timeline,
                "hf_token": HF_TOKEN
            }
        }
        
        resp = http_requests.post(url, headers=headers, json=payload)
        resp_data = resp.json()
        job_id = resp_data.get("id")
        
        if job_id:
            task["status"] = "transcribing"
            task["progress"] = 10
            task["job_id"] = job_id
            
            threading.Thread(target=poll_job, args=(job_id, task_id), daemon=True).start()
            print(f"🚀 Serverless Transcription Job Started: {job_id} for {task_id}")
            return {"status": "started", "job_id": job_id}
        else:
            return {"status": "error", "error": f"Failed to start job: {resp_data}"}

    except Exception as e:
        task["status"] = "error"
        task["error"] = str(e)
        return {"status": "error", "error": str(e)}



# ═══════════════════════════════════════════
#  API ENDPOINTS
# ═══════════════════════════════════════════

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Save file locally and begin S3 upload in background."""
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    task_id = file.filename

    # Check if we already have a transcription for this file
    json_path = file_path.with_suffix(".json")
    if json_path.exists():
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                transcriptions[task_id] = data
                print(f"📎 Found existing transcription for {task_id}")
                return {"task_id": task_id}
        except:
            pass

    transcriptions[task_id] = {
        "filename": file.filename,
        "status": "uploading",
        "progress": 0,
        "result": [],
    }

    # Start S3 upload in background
    t = threading.Thread(target=upload_to_s3, args=(file_path, task_id), daemon=True)
    t.start()

    return {"task_id": task_id}


@app.get("/check/{filename}")
async def check_transcription(filename: str):
    """Check if a transcription JSON already exists for this audio file."""
    # Check in-memory first
    if filename in transcriptions and transcriptions[filename].get("status") == "completed":
        return transcriptions[filename]

    # Check on disk
    json_path = UPLOAD_DIR / Path(filename).with_suffix(".json")
    if json_path.exists():
        try:
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                transcriptions[filename] = data
                return data
        except:
            pass

    return {"status": "not_found"}


@app.get("/status/{task_id}")
async def get_status(task_id: str):
    """Return task status."""
    task = transcriptions.get(task_id)
    if not task:
        return {"status": "not_found"}
    return task


@app.get("/audio/{filename}")
async def get_audio(filename: str):
    """Serve audio file for the local player."""
    return FileResponse(UPLOAD_DIR / filename)


@app.get("/download/{filename}")
async def download_file(filename: str):
    """Download .md or .docx result files."""
    path = UPLOAD_DIR / filename
    if path.exists():
        media_type = "text/markdown"
        if filename.endswith(".docx"):
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(path, media_type=media_type, filename=filename)
    return {"error": "File not found"}


@app.post("/save/{task_id}")
async def save_files(task_id: str):
    """Generate and save .md and .docx from the current transcription state."""
    if task_id not in transcriptions:
        return {"error": "Task not found"}

    task = transcriptions[task_id]
    if not task.get("result"):
        return {"error": "No transcription data to save"}

    regenerate_files(task_id)
    return {
        "status": "saved",
        "md_path": Path(task["filename"]).with_suffix(".md").name,
        "docx_path": Path(task["filename"]).with_suffix(".docx").name,
    }


class UpdateSpeakerRequest(BaseModel):
    task_id: str
    segment_index: int
    speaker_name: str


@app.post("/update_speaker")
async def update_speaker(req: UpdateSpeakerRequest):
    """Bulk rename a speaker across all segments."""
    if req.task_id in transcriptions:
        task = transcriptions[req.task_id]
        if 0 <= req.segment_index < len(task["result"]):
            old_name = task["result"][req.segment_index]["speaker"]
            new_name = req.speaker_name

            for seg in task["result"]:
                if seg["speaker"] == old_name:
                    seg["speaker"] = new_name

            regenerate_files(req.task_id)
            return {"status": "success"}

    return {"status": "error", "message": "Task or segment not found"}


@app.get("/list")
async def list_transcriptions():
    """List all available transcriptions (for a file picker)."""
    results = []
    for json_file in UPLOAD_DIR.glob("*.json"):
        try:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
                results.append({
                    "filename": data.get("filename"),
                    "segments": len(data.get("result", [])),
                    "status": data.get("status", "unknown"),
                })
        except:
            pass
    return results


# ═══════════════════════════════════════════
#  FULL POD AUTOMATION (SSH + GQL)
# ═══════════════════════════════════════════

class PodConfigRequest(BaseModel):
    ip: str = ""
    ssh_port: int = 22
    pod_id: str = ""
    endpoint_id: str = ""
    key_path: str = None

@app.post("/update-pod-config")
async def update_pod_config(req: PodConfigRequest):
    """Save Pod metadata and Serverless Endpoint ID to the current session and .env."""
    global RUNPOD_POD_ID, RUNPOD_ENDPOINT_ID
    
    if req.ip: pod_config["ip"] = req.ip
    if req.ssh_port: pod_config["ssh_port"] = req.ssh_port
    if req.pod_id: RUNPOD_POD_ID = req.pod_id
    if req.endpoint_id: RUNPOD_ENDPOINT_ID = req.endpoint_id
    if req.key_path: pod_config["key_path"] = req.key_path
    
    # Cleanly update .env
    env_lines = []
    if os.path.exists(".env"):
        with open(".env", "r") as f:
            env_lines = f.readlines()
    
    # Filter out old keys
    keys_to_remove = ["POD_IP=", "POD_SSH_PORT=", "RUNPOD_POD_ID=", "POD_KEY_PATH=", "RUNPOD_ENDPOINT_ID="]
    env_lines = [l for l in env_lines if not any(k in l for k in keys_to_remove)]
    
    # Add new values (preserving some defaults if not provided in req but present in session)
    if pod_config["ip"]: env_lines.append(f"POD_IP={pod_config['ip']}\n")
    if pod_config["ssh_port"]: env_lines.append(f"POD_SSH_PORT={pod_config['ssh_port']}\n")
    if RUNPOD_POD_ID: env_lines.append(f"RUNPOD_POD_ID={RUNPOD_POD_ID}\n")
    if RUNPOD_ENDPOINT_ID: env_lines.append(f"RUNPOD_ENDPOINT_ID={RUNPOD_ENDPOINT_ID}\n")
    if pod_config["key_path"]: env_lines.append(f"POD_KEY_PATH={pod_config['key_path']}\n")
    
    with open(".env", "w") as f:
        f.writelines(env_lines)
    
    print(f"📡 Config Updated. Endpoint: {RUNPOD_ENDPOINT_ID}")
    return {"status": "updated"}

@app.get("/get-pod-config")
async def get_pod_config():
    """Return the current active config."""
    return {
        "ip": pod_config.get("ip"),
        "ssh_port": pod_config.get("ssh_port"),
        "pod_id": RUNPOD_POD_ID,
        "endpoint_id": RUNPOD_ENDPOINT_ID,
        "key_path": pod_config.get("key_path")
    }

def get_ssh_client():
    """Create an SSH client for the Pod, auto-detecting the SSH key."""
    ssh = paramiko.SSHClient()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())

    # Try key locations in order of preference
    key_candidates = [
        pod_config.get("key_path", "runpod"),           # Project key
        os.path.expanduser("~/.ssh/id_ed25519"),        # Default Ed25519
        os.path.expanduser("~/.ssh/id_rsa"),            # Default RSA
    ]

    last_error = None
    for key_path in key_candidates:
        if not os.path.exists(key_path):
            continue
        try:
            # Try Ed25519 first, then RSA
            for KeyClass in [paramiko.Ed25519Key, paramiko.RSAKey]:
                try:
                    key = KeyClass.from_private_key_file(key_path)
                    ssh.connect(pod_config["ip"], port=pod_config["ssh_port"], username="root", pkey=key, timeout=10)
                    print(f"✅ SSH connected via {key_path}")
                    return ssh
                except (paramiko.ssh_exception.SSHException, ValueError):
                    continue
        except Exception as e:
            last_error = e
            continue

    raise Exception(f"Could not connect with any SSH key. Last error: {last_error}")

@app.post("/setup-pod")
async def setup_pod():
    """Deploy code to Pod and run the setup script."""
    try:
        ssh = get_ssh_client()
        
        # 1. Archive current project (excluding uploads, cache, .git)
        archive_path = "project.tar.gz"
        print("📦 Creating project archive...")
        with tarfile.open(archive_path, "w:gz") as tar:
            for item in os.listdir("."):
                if item in ["uploads", "cache", ".git", "__pycache__", "legacy", "logs", "project.tar.gz"]:
                    continue
                tar.add(item)
        
        # 2. Upload to Pod
        print("📤 Uploading project to Pod...")
        with SCPClient(ssh.get_transport()) as scp:
            scp.put(archive_path, "/workspace/project.tar.gz")
        
        # 3. Extract and Setup
        print("🛠️ Running setup on Pod...")
        commands = [
            "cd /workspace && mkdir -p transcriber && tar -xzf project.tar.gz -C transcriber",
            "cd /workspace/transcriber && mkdir -p uploads cache",
            f"echo 'RUNPOD_ACCESS_KEY={os.getenv('RUNPOD_ACCESS_KEY')}' > /workspace/transcriber/.env",
            f"echo 'RUNPOD_SECRET_KEY={os.getenv('RUNPOD_SECRET_KEY')}' >> /workspace/transcriber/.env",
            "cd /workspace/transcriber && bash setup_runpod.sh > worker.log 2>&1"
        ]
        
        full_cmd = " && ".join(commands)
        stdin, stdout, stderr = ssh.exec_command(full_cmd)
        
        # We don't wait for setup to finish in the request, we return 200 and let it run
        # but for this simplified version, let's at least start it.
        return {"status": "setup_started", "message": "Deploying and installing dependencies..."}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/start-transcription")
async def start_transcription():
    """Start the remote worker in a screen session."""
    try:
        ssh = get_ssh_client()
        # Start in a screen session so it persists
        cmd = "screen -dmS transcriber bash -c 'cd /workspace/transcriber && python remote_worker.py --watch > worker.log 2>&1'"
        ssh.exec_command(cmd)
        return {"status": "started", "message": "Worker started on Pod."}
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.get("/pod-logs")
async def get_pod_logs():
    """Fetch the latest logs from the Pod worker."""
    try:
        ssh = get_ssh_client()
        stdin, stdout, stderr = ssh.exec_command("tail -n 50 /workspace/transcriber/worker.log")
        logs = stdout.read().decode("utf-8")
        return {"logs": logs}
    except Exception as e:
        return {"logs": f"Error fetching logs: {str(e)}"}

def runpod_gql(query):
    """Send a GraphQL request to RunPod API."""
    headers = {"Content-Type": "application/json", "api-key": RUNPOD_API_KEY}
    resp = http_requests.post(RUNPOD_GQL, json={"query": query}, headers=headers)
    return resp.json()


@app.post("/start-pod")
async def start_pod():
    """Start the RunPod GPU Pod remotely."""
    if not RUNPOD_POD_ID:
        return {"error": "RUNPOD_POD_ID not set in .env"}
    query = f'mutation {{ podResume(input: {{ podId: "{RUNPOD_POD_ID}", gpuCount: 1 }}) {{ id desiredStatus }} }}'
    result = runpod_gql(query)
    print(f"🚀 Pod start requested: {result}")
    return {"status": "starting", "result": result}


@app.post("/stop-pod")
async def stop_pod():
    """Stop the RunPod GPU Pod remotely."""
    if not RUNPOD_POD_ID:
        return {"error": "RUNPOD_POD_ID not set in .env"}
    query = f'mutation {{ podStop(input: {{ podId: "{RUNPOD_POD_ID}" }}) {{ id desiredStatus }} }}'
    result = runpod_gql(query)
    print(f"⏹️ Pod stop requested: {result}")
    return {"status": "stopping", "result": result}


@app.get("/pod-status")
async def pod_status():
    """Check the current status of the RunPod Pod."""
    if not RUNPOD_POD_ID:
        return {"status": "not_configured"}
    try:
        query = f'{{ pod(input: {{ podId: "{RUNPOD_POD_ID}" }}) {{ id name desiredStatus runtime {{ uptimeInSeconds }} }} }}'
        result = runpod_gql(query)
        pod = result.get("data", {}).get("pod") or {}
        if not pod:
            return {"status": "NOT_FOUND"}
        return {
            "status": pod.get("desiredStatus", "UNKNOWN"),
            "name": pod.get("name", ""),
            "uptime": pod.get("runtime", {}).get("uptimeInSeconds") if pod.get("runtime") else None,
        }
    except Exception as e:
        return {"status": "ERROR", "message": str(e)}


@app.post("/sync-now")
async def sync_now():
    """Manually trigger a cloud sync check."""
    download_results_from_s3()
    return {"status": "synced", "tasks": len(transcriptions)}


# ─── Static Files & Startup ───
app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8001))
    print(f"🚀 Transcriber Pro (Local Dashboard) starting on port {port}...")
    uvicorn.run(app, host="0.0.0.0", port=port)
