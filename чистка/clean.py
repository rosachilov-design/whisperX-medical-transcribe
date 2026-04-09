import docx
import re

def clean_doc(input_path, output_path):
    doc = docx.Document(input_path)
    new_doc = docx.Document()

    # corrections
    corrections = {
        r"глюкофашлон": "Глюкофаж Лонг",
        r"глюкопашленку": "Глюкофаж Лонг",
        r"глюкопашлон": "Глюкофаж Лонг",
        r"глюкофаршлонг": "Глюкофаж Лонг",
        r"глюкозит": "гликлазид",
        r"гликозид": "гликлазид",
        r"ИБСГ": "ИБС с АГ",
        r"хосеенщики": "ХСН-щики",
        r"АЙ-20\.8": "I20.8",
        r"АЙ-11\.9": "I11.9",
        r"Ай-11\.9": "I11.9",
        r"АЙ-48": "I48",
        r"Ай-48": "I48",
        r"Ай-40": "I40",
        r"АК-6": "АКШ",
        r"бисепролол": "бисопролол",
        r"младипин": "амлодипин",
        r"не бивалол": "небиволол",
        r"небе волос": "небиволол",
        r"Комкорн": "Конкор",
        r"пристилол": "Престилол",
        r"простилол": "Престилол",
        r"метаформин": "метформин",
        r"инфлины": "инсулины",
        r"хобликов": "ХОБЛиков",
        r"предебит": "преддиабет",
        r"Нолипреллы": "Нолипрелы",
        r"Азон": "Озон",
        r"Назвезда": "Северная Звезда",
        r"Нули Прел": "Нолипрел",
        r"Аби Форте": "А Би-форте",
        r"доскречена": "доступна",
        r"в инвитро": "в Инвитро",
        r"децентрализацию": "диспансеризацию"
    }

    def apply_corrections(text):
        for w, rep in corrections.items():
            text = re.sub(rf'(?i)\b{w}\b', rep, text)
        return text

    blocks = []
    current_speaker = None
    current_text = []

    pattern = re.compile(r'^(?:\[\d{2}:\d{2}(?::\d{2})?\]\s*)?(SPEAKER_00|SPEAKER_01|Unknown)\s*:\s*(.*)$')

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        
        match = pattern.match(txt)
        if match:
            speaker = match.group(1)
            content = match.group(2)
            if speaker == current_speaker:
                current_text.append(content)
            else:
                if current_speaker is not None:
                    blocks.append((current_speaker, " ".join(current_text)))
                current_speaker = speaker
                current_text = [content]
        else:
            if current_speaker is not None:
                current_text.append(txt)
            else:
                blocks.append((None, txt))

    if current_speaker is not None:
        blocks.append((current_speaker, " ".join(current_text)))

    for speaker, text in blocks:
        text = apply_corrections(text)
        
        p = new_doc.add_paragraph()
        if speaker == "SPEAKER_00":
            r = p.add_run(f"Интервьюер: {text}")
            r.bold = True
        elif speaker == "SPEAKER_01":
            r_sp = p.add_run("Респондент: ")
            r_sp.bold = True
            p.add_run(text)
        elif speaker == "Unknown":
            r_sp = p.add_run("Unknown: ")
            r_sp.bold = True
            p.add_run(text)
        else:
            p.add_run(text)

    new_doc.save(output_path)

if __name__ == "__main__":
    clean_doc("uploads/interview3.docx", "uploads/Interview3_Cleaned.docx")
