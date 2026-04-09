
import docx
import re

def process_transcript(input_path, output_path, vocab_path):
    doc = docx.Document(input_path)
    
    # Speaker mapping based on our analysis
    # SPEAKER_01 -> Интервьюер
    # SPEAKER_00 -> Респондент
    speaker_map = {
        "SPEAKER_01": "Интервьюер",
        "SPEAKER_00": "Респондент"
    }

    processed_blocks = []
    current_speaker = None
    current_text = ""
    current_time = ""

    # Regexp to parse: [00:10] SPEAKER_00: text
    pattern = re.compile(r"\[(\d{2}:\d{2}(?::\d{2})?)\] (SPEAKER_\d{2}): (.*)")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        match = pattern.match(text)
        if match:
            time_str, speaker_id, content = match.groups()
            speaker_name = speaker_map.get(speaker_id, speaker_id)
            
            if speaker_name == current_speaker:
                # Merge with previous
                current_text += " " + content
            else:
                # Save previous block
                if current_speaker:
                    processed_blocks.append({
                        "speaker": current_speaker,
                        "time": current_time,
                        "text": current_text.strip()
                    })
                # Start new block
                current_speaker = speaker_name
                current_time = time_str
                current_text = content
        else:
            # If it doesn't match the pattern but there's text, append to current block or handle as raw
            if current_speaker:
                current_text += " " + text
            else:
                # Likely header or something else
                processed_blocks.append({
                    "speaker": None,
                    "time": None,
                    "text": text
                })

    # Add the last block
    if current_speaker:
        processed_blocks.append({
            "speaker": current_speaker,
            "time": current_time,
            "text": current_text.strip()
        })

    # Create new document
    new_doc = docx.Document()
    
    medical_terms = set()
    
    # Basic spelling fix (very basic)
    def fix_spelling(text):
        # We'll rely on a list of known medical terms or just provide the text as is 
        # but I can do some fixes for obvious stuff if I find any.
        # For now, let's keep it mostly as is and I'll do a pass over the text if needed.
        # Actually, let's just use the text and I will do the extraction.
        return text

    for block in processed_blocks:
        p = new_doc.add_paragraph()
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0
        
        speaker = block["speaker"]
        time = block["time"]
        text = block["text"]
        
        if speaker == "Интервьюер":
            # Bold all words
            run = p.add_run(f"Интервьюер: [{time}] {text}")
            run.bold = True
        elif speaker == "Респондент":
            # Only "Респондент" bold
            run_tag = p.add_run("Респондент")
            run_tag.bold = True
            run_rest = p.add_run(f": [{time}] {text}")
        else:
            p.add_run(text)
            
        # Collect potential medical terms (very rudimentary - will improve)
        # In a real scenario I would use a list or LLM to extract.
        # Since I am the LLM, I will do this by processing the text after reading it.
    
    new_doc.save(output_path)
    return processed_blocks

# List of terms to look for based on initial scan
initial_terms = [
    "глиптины", "аллоглиптин", "симовик", "оземпик", "инсулин", 
    "сахарный диабет", "углеводный обмен", "гипертония", "диспансеризация",
    "терапевт", "эндокринолог", "сахороснижающая терапия", "глюкоза"
]

if __name__ == "__main__":
    input_file = r"c:\Users\halfo\OneDrive\Documents\GitHub\whisperX-medical-transcribe\uploads\ГИ4_Тер_гос_Спб_КораблеваНВ_20.02.docx"
    output_file = r"c:\Users\halfo\OneDrive\Documents\GitHub\whisperX-medical-transcribe\uploads\ГИ4_Тер_гос_Спб_КораблеваНВ_20.02_fixed.docx"
    
    blocks = process_transcript(input_file, output_file, None)
    print(f"Processed {len(blocks)} blocks.")
    
    # Extracting vocab from blocks
    all_text = " ".join([b["text"] for b in blocks])
    with open("transcript_text.txt", "w", encoding="utf-8") as f:
        f.write(all_text)
    print(f"Text saved to transcript_text.txt. Total length: {len(all_text)}")
