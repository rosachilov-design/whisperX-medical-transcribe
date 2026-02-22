from fastapi import FastAPI, UploadFile, File, BackgroundTasks, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from faster_whisper import WhisperModel
import os
import torch
from pathlib import Path
import json
import asyncio
import time
import subprocess
import math
import re
import threading
from datetime import datetime


from docx import Document
import torchaudio
if not hasattr(torchaudio, "list_audio_backends"):
    torchaudio.list_audio_backends = lambda: ["soundfile"]
from pyannote.audio import Pipeline



app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
CACHE_DIR = Path("cache")
CACHE_DIR.mkdir(exist_ok=True)

# Transcriptions storage
transcriptions = {}

def log_info(message):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] [VERSION 2.0.1] {message}")

def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    if hours > 0:
        return f"{hours:02}:{minutes:02}:{secs:02}"
    return f"{minutes:02}:{secs:02}"

def get_duration(file_path):
    cmd = ["ffprobe", "-v", "error", "-show_entries", "format=duration", "-of", "default=noprint_wrappers=1:nokey=1", str(file_path)]
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True)
    return float(result.stdout.strip())

# Load models once
device = "cuda" if torch.cuda.is_available() else "cpu"

# Set UTF-8 for Windows output redirection
import sys
import io
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

model_lock = threading.Lock()

log_info(f"Loading faster-whisper model on {device}...")
compute_type = "float16" if device == "cuda" else "int8"
model = WhisperModel("turbo", device=device, compute_type=compute_type)
log_info("faster-whisper loaded.")

HF_TOKEN = os.getenv("HF_TOKEN")
print("Loading Diarization Pipeline...")
try:
    diarization_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        token=HF_TOKEN
    )
    if diarization_pipeline:
        diarization_pipeline.to(torch.device(device))
        log_info("Diarization Pipeline loaded.")
    else:
        log_info("Failed to load Diarization Pipeline (check token/access).")
except Exception as e:
    print(f"Error loading diarization pipeline: {e}")
    diarization_pipeline = None

def run_diarization(file_path: Path, task_id: str = None):
    if not diarization_pipeline:
        return []
    
    # Cache diarization results to a JSON file to avoid 40min re-runs
    cache_file = CACHE_DIR / f"{file_path.stem}_diarize.json"
    if cache_file.exists():
        log_info(f"Loading cached diarization for {file_path.name}...")
        with open(cache_file, 'r') as f:
            return json.load(f)

    log_info(f"Running diarization on {file_path.name}...")
    wav_path = CACHE_DIR / f"{file_path.stem}_diarize.wav"
    cmd = [
        "ffmpeg", "-y", "-i", str(file_path),
        "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(wav_path)
    ]
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    
    duration = get_duration(file_path)
    log_info(f"Audio converted. Duration: {duration:.2f}s. Loading with soundfile...")
    
    import soundfile as sf
    import numpy as np
    data, sample_rate = sf.read(str(wav_path), dtype='float32')
    if data.ndim == 1:
        data = data[np.newaxis, :]
    else:
        data = data.T
    waveform = torch.from_numpy(data)
    audio_input = {"waveform": waveform, "sample_rate": sample_rate}
    
    log_info("Starting pyannote pipeline with progress tracking...")
    
    def hook(step_name, step_artifact, file=None, **kwargs):
        if task_id and task_id in transcriptions:
            try:
                # Add a quick debug print to see if pyannote is moving or stuck
                print(f"[Pyannote] Step: {step_name}")
                from pyannote.core import Segment
                if isinstance(step_artifact, Segment):
                    # Clamp progress to 99% during diarization phase
                    p = min(99, int((step_artifact.end / duration) * 100))
                    if p > transcriptions[task_id].get("progress", 0):
                        transcriptions[task_id]["progress"] = p
            except Exception:
                pass

    with model_lock:
        diarize_output = diarization_pipeline(audio_input, min_speakers=2, hook=hook)
    log_info("Diarization complete.")
    
    if hasattr(diarize_output, 'speaker_diarization'):
        annotation = diarize_output.speaker_diarization
    else:
        annotation = diarize_output
    
    timeline = []
    for turn, _, speaker in annotation.itertracks(yield_label=True):
        timeline.append({
            "start": turn.start,
            "end": turn.end,
            "speaker": speaker
        })
        
    # Save to cache
    with open(cache_file, 'w') as f:
        json.dump(timeline, f)

    if wav_path.exists():
        os.remove(wav_path)
        
    return timeline

def clean_hallucinations(text: str) -> str:
    """Remove common Russian Whisper hallucinations like 'Subtitle editor', etc.
    Uses non-greedy matching to avoid eating actual speech that follows."""
    # We look for the keyword and then a typical name structure (Initials + Surname)
    # or just the keyword itself if isolated.
    hallucination_patterns = [
        r'\bРедактор субтитров\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bКорректор\s+([А-ЯA-Z]\.?\s*){1,2}[А-ЯA-Z][а-яa-z]+',
        r'\bСубтитры\s*:\s*[^\.]+',
        r'\bПеревод\s*:\s*[^\.]+',
        r'\bОзвучка\s*:\s*[^\.]+',
        r'\bРедактор субтитров\b',
        r'\bКорректор\b',
        r'\b(Все права защищены|Продолжение следует|Ставьте лайки|Подписывайтесь на канал)\b',
    ]
    
    cleaned = text
    for pattern in hallucination_patterns:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)
    
    # Final cleanup of double spaces or leftover dangling punctuation
    cleaned = re.sub(r'\s+', ' ', cleaned)
    return cleaned.strip()


def get_speaker_for_word(timeline, word_start, word_end):
    """Assign a speaker to a single word using its timestamp against the diarization timeline."""
    if not timeline:
        return "Unknown"
    
    best_speaker = None
    best_overlap = 0
    
    for entry in timeline:
        overlap_start = max(word_start, entry["start"])
        overlap_end = min(word_end, entry["end"])
        
        if overlap_end > overlap_start:
            overlap = overlap_end - overlap_start
            if overlap > best_overlap:
                best_overlap = overlap
                best_speaker = entry["speaker"]
    
    if best_speaker:
        return best_speaker
    
    # Fallback: find nearest diarization turn
    mid = (word_start + word_end) / 2
    nearest = min(timeline, key=lambda s: min(abs(s["start"] - mid), abs(s["end"] - mid)))
    return nearest["speaker"]


def generate_docx(task_id):
    if task_id not in transcriptions:
        return None
    
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]
    docx_file_path = file_path.with_suffix(".docx")
    
    doc = Document()
    doc.add_heading(f"Transcription: {task['filename']}", 0)
    
    for seg in task["result"]:
        p = doc.add_paragraph()
        ts_run = p.add_run(f"[{seg['timestamp']}] {seg['speaker']}: ")
        ts_run.bold = True
        p.add_run(seg['text'])
        
    doc.save(docx_file_path)
    return docx_file_path.name


def self_group_words(speaker_words, speaker_map, speaker_counter):
    """Group words by speaker for live display during transcription."""
    segments = []
    if not speaker_words:
        return segments
    
    cur_raw = speaker_words[0]["speaker_raw"]
    cur_start = speaker_words[0]["start"]
    cur_words = [speaker_words[0]["word"]]
    
    for sw in speaker_words[1:]:
        if sw["speaker_raw"] == cur_raw:
            cur_words.append(sw["word"])
        else:
            if cur_raw not in speaker_map:
                speaker_counter += 1
                speaker_map[cur_raw] = f"Speaker {speaker_counter}"
            text = clean_hallucinations(" ".join(cur_words))
            if text:
                segments.append({
                    "start": cur_start,
                    "timestamp": format_timestamp(cur_start),
                    "text": re.sub(r'\s+', ' ', text),
                    "speaker": speaker_map[cur_raw]
                })
            cur_raw = sw["speaker_raw"]
            cur_start = sw["start"]
            cur_words = [sw["word"]]
    
    if cur_words:
        if cur_raw not in speaker_map:
            speaker_counter += 1
            speaker_map[cur_raw] = f"Speaker {speaker_counter}"
        text = clean_hallucinations(" ".join(cur_words))
        if text:
            segments.append({
                "start": cur_start,
                "timestamp": format_timestamp(cur_start),
                "text": re.sub(r'\s+', ' ', text),
                "speaker": speaker_map[cur_raw]
            })
    
    # Merge adjacent same-speaker
    merged = []
    for seg in segments:
        if merged and seg["speaker"] == merged[-1]["speaker"]:
            merged[-1]["text"] += " " + seg["text"]
        else:
            merged.append(seg.copy())
    
    return merged

def run_diarize_task(file_path: Path, task_id: str):
    try:
        transcriptions[task_id]["status"] = "diarizing"
        transcriptions[task_id]["progress"] = 0
        timeline = run_diarization(file_path, task_id=task_id)
        
        # Cache timeline in transcription object for the next step
        transcriptions[task_id]["timeline"] = timeline
        transcriptions[task_id]["status"] = "diarization_complete"
        transcriptions[task_id]["progress"] = 100
        log_info(f"Diarization complete for {task_id}.")
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error in diarization task: {e}")
        transcriptions[task_id]["status"] = "error"
        transcriptions[task_id]["error"] = str(e)


def run_transcribe_task(file_path: Path, task_id: str):
    try:
        timeline = transcriptions[task_id].get("timeline", [])
        
        # Phase 2: Silence-aware chunked transcription
        transcriptions[task_id]["status"] = "transcribing"
        transcriptions[task_id]["progress"] = 10
        
        # Build natural chunks based on diarization timeline
        # We group turns into blocks of roughly 30-45 seconds, splitting only at gaps
        natural_chunks = []
        if timeline:
            current_chunk_turns = []
            chunk_start_time = timeline[0]["start"]
            
            for i, entry in enumerate(timeline):
                current_chunk_turns.append(entry)
                
                # Check if we should close this chunk
                elapsed = entry["end"] - chunk_start_time
                
                # If we've reached ~30s, look for a gap or just split at the end of this turn
                is_last = (i == len(timeline) - 1)
                
                if elapsed >= 30 or is_last:
                    # Found a natural boundary (end of a speaker turn)
                    natural_chunks.append({
                        "start": chunk_start_time,
                        "end": entry["end"],
                        "turns": current_chunk_turns
                    })
                    
                    if not is_last:
                        current_chunk_turns = []
                        chunk_start_time = timeline[i+1]["start"]
        
        if not natural_chunks:
            # Fallback if no diarization results
            duration = get_duration(file_path)
            natural_chunks = [{"start": i*30, "end": min((i+1)*30, duration)} for i in range(math.ceil(duration/30))]

        # Decode entire file to WAV once to prevent ffmpeg from scanning the large source file repeatedly
        full_wav_path = CACHE_DIR / f"{task_id}_full.wav"
        if not full_wav_path.exists():
            log_info(f"Decoding full audio to WAV for fast extraction...")
            subprocess.run([
                "ffmpeg", "-y", "-i", str(file_path), "-ar", "16000", "-ac", "1", "-c:a", "pcm_s16le", str(full_wav_path)
            ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        speaker_map = {}
        speaker_counter = 0
        all_speaker_words = []
        
        for i, chunk in enumerate(natural_chunks):
            actual_start = chunk["start"]
            actual_end = chunk["end"]
            
            # 10 seconds of overlapping context to improve boundary precision
            pad = 10.0
            start_time = max(0, actual_start - pad)
            end_time = actual_end + pad
            duration_s = end_time - start_time
            
            if duration_s <= 0:
                continue
                
            chunk_path = CACHE_DIR / f"{task_id}_chunk_{i}.wav"
            
            # Extract chunk efficiently from the uncompressed WAV and properly rewrite headers
            cmd = [
                "ffmpeg", "-y", "-ss", str(start_time), "-t", str(duration_s),
                "-i", str(full_wav_path), "-c:a", "pcm_s16le", str(chunk_path)
            ]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            if not chunk_path.exists():
                continue
            
            # Context prompt
            previous_context = ""
            if all_speaker_words:
                last_words = [sw["word"] for sw in all_speaker_words[-50:]]
                previous_context = " ".join(last_words)
            
            log_info(f"Transcribing natural chunk {i+1}/{len(natural_chunks)} ({start_time:.2f}s - {end_time:.2f}s)")
            
            try:
                with model_lock:
                    # Run fast transcription using faster-whisper
                    segments, info = model.transcribe(
                        str(chunk_path),
                        language="ru",
                        word_timestamps=True,
                        condition_on_previous_text=True,
                        initial_prompt=previous_context if previous_context else "Это аудиозапись беседы или интервью.",
                    )
                    # Force evaluation of the generator and list conversion inside the lock
                    segments = list(segments)

                for segment in segments:
                    words = getattr(segment, "words", [])
                    if not words:
                        text = re.sub(r'\[.*?\]', '', segment.text).strip()
                        if text:
                            abs_start = segment.start + start_time
                            abs_end = segment.end + start_time
                            midpoint = (abs_start + abs_end) / 2.0
                            if actual_start <= midpoint <= actual_end:
                                raw = get_speaker_for_word(timeline, abs_start, abs_end)
                                all_speaker_words.append({
                                    "word": text, "start": abs_start, "end": abs_end, "speaker_raw": raw
                                })
                        continue
                    
                    for w in words:
                        word_text = getattr(w, "word", "").strip()
                        if not word_text: continue
                        abs_start = w.start + start_time
                        abs_end = w.end + start_time
                        midpoint = (abs_start + abs_end) / 2.0
                        if actual_start <= midpoint <= actual_end:
                            raw = get_speaker_for_word(timeline, abs_start, abs_end)
                            all_speaker_words.append({
                                "word": word_text, "start": abs_start, "end": abs_end, "speaker_raw": raw
                            })
            except Exception as chunk_err:
                if "cublas" in str(chunk_err).lower() or "cudnn" in str(chunk_err).lower() or getattr(chunk_err, "message", "") == "Library cublas64_12.dll is not found or cannot be loaded":
                    log_info(f"WARNING: Chunk {i+1} faster-whisper DLL error ({chunk_err}), falling back to openai-whisper...")
                    import whisper
                    global fallback_model
                    if "fallback_model" not in globals():
                        log_info(f"Loading robust whisper fallback model on {device}...")
                        fallback_model = whisper.load_model("turbo", device=device)
                    
                    try:
                        with model_lock:
                            result = fallback_model.transcribe(
                                str(chunk_path),
                                language="ru",
                                verbose=False,
                                fp16=True, # Try FP16 first
                                word_timestamps=True,
                                condition_on_previous_text=True,
                                initial_prompt=previous_context if previous_context else "Это аудиозапись беседы или интервью.",
                            )
                    except Exception as e2:
                        log_info(f"FP16 fallback failing ({e2}), retrying FP32...")
                        with model_lock:
                            result = fallback_model.transcribe(
                                str(chunk_path), language="ru", verbose=False, fp16=False, word_timestamps=True,
                                condition_on_previous_text=True, initial_prompt=previous_context if previous_context else "Это аудиозапись беседы или интервью."
                            )
                    
                    for segment in result["segments"]:
                        words = segment.get("words", [])
                        if not words:
                            text = re.sub(r'\[.*?\]', '', segment["text"]).strip()
                            if text:
                                abs_start = segment["start"] + start_time
                                abs_end = segment["end"] + start_time
                                midpoint = (abs_start + abs_end) / 2.0
                                if actual_start <= midpoint <= actual_end:
                                    raw = get_speaker_for_word(timeline, abs_start, abs_end)
                                    all_speaker_words.append({
                                        "word": text, "start": abs_start, "end": abs_end, "speaker_raw": raw
                                    })
                            continue
                        
                        for w in words:
                            word_text = w.get("word", "").strip()
                            if not word_text: continue
                            abs_start = w["start"] + start_time
                            abs_end = w["end"] + start_time
                            midpoint = (abs_start + abs_end) / 2.0
                            if actual_start <= midpoint <= actual_end:
                                raw = get_speaker_for_word(timeline, abs_start, abs_end)
                                all_speaker_words.append({
                                    "word": word_text, "start": abs_start, "end": abs_end, "speaker_raw": raw
                                })
                else:
                    log_info(f"WARNING: Chunk {i+1} failed completely ({chunk_err}), skipping...")
            finally:
                if chunk_path.exists():
                    os.remove(chunk_path)
            
            # Progress tracking
            transcriptions[task_id]["progress"] = 10 + int(((i + 1) / len(natural_chunks)) * 80)
            
            # Live result
            live_segments = self_group_words(all_speaker_words, speaker_map, speaker_counter)
            if live_segments:
                speaker_counter = max(speaker_counter, len(speaker_map))
                transcriptions[task_id]["result"] = live_segments
        
        log_info(f"Assigned speakers to {len(all_speaker_words)} words across {len(natural_chunks)} natural chunks.")
        
        if full_wav_path.exists():
            os.remove(full_wav_path)
            
        # Phase 3: Final grouping
        transcriptions[task_id]["status"] = "aligning"
        transcriptions[task_id]["progress"] = 95

        
        # Group consecutive words by the same speaker into segments
        final_segments = []
        if all_speaker_words:
            current_speaker_raw = all_speaker_words[0]["speaker_raw"]
            current_start = all_speaker_words[0]["start"]
            current_words = [all_speaker_words[0]["word"]]
            
            for sw in all_speaker_words[1:]:
                if sw["speaker_raw"] == current_speaker_raw:
                    current_words.append(sw["word"])
                else:
                    # Flush current segment
                    if current_speaker_raw not in speaker_map:
                        speaker_counter += 1
                        speaker_map[current_speaker_raw] = f"Speaker {speaker_counter}"
                    
                    text = clean_hallucinations(" ".join(current_words))
                    text = re.sub(r'\s+', ' ', text)
                    if text:
                        final_segments.append({
                            "start": current_start,
                            "timestamp": format_timestamp(current_start),
                            "text": text,
                            "speaker": speaker_map[current_speaker_raw]
                        })
                    
                    current_speaker_raw = sw["speaker_raw"]
                    current_start = sw["start"]
                    current_words = [sw["word"]]
            
            # Flush last segment
            if current_words:
                if current_speaker_raw not in speaker_map:
                    speaker_counter += 1
                    speaker_map[current_speaker_raw] = f"Speaker {speaker_counter}"
                text = clean_hallucinations(" ".join(current_words))
                text = re.sub(r'\s+', ' ', text)
                if text:
                    final_segments.append({
                        "start": current_start,
                        "timestamp": format_timestamp(current_start),
                        "text": text,
                        "speaker": speaker_map[current_speaker_raw]
                    })
        
        # Merge adjacent same-speaker segments
        smoothed = []
        for seg in final_segments:
            if smoothed and seg["speaker"] == smoothed[-1]["speaker"]:
                smoothed[-1]["text"] += " " + seg["text"]
            else:
                smoothed.append(seg.copy())
        
        transcriptions[task_id]["result"] = smoothed


        # Final save
        md_content = f"# Transcription: {file_path.name}\n\n"
        for seg in smoothed:
            md_content += f"**[{seg['timestamp']}] {seg['speaker']}:** {seg['text']}\n\n"
        
        md_file_path = file_path.with_suffix(".md")
        with open(md_file_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        
        # Generate DOCX
        docx_name = generate_docx(task_id)
            
        transcriptions[task_id].update({
            "status": "completed",
            "progress": 100,
            "md_path": str(md_file_path.name),
            "docx_path": docx_name
        })

        log_info(f"Transcription complete: {len(smoothed)} speaker turns.")
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error in live transcription: {e}")
        transcriptions[task_id]["status"] = "error"
        transcriptions[task_id]["error"] = str(e)

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_path = UPLOAD_DIR / file.filename
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())
    
    task_id = file.filename
    log_info(f"Upload received: {task_id}")
    
    status = "uploaded"
    progress = 0
    timeline = []
    result = []
    
    # Auto-detect cache
    stem = Path(file.filename).stem
    cache_file = CACHE_DIR / f"{stem}_diarize.json"
    full_transcription_file = UPLOAD_DIR / f"{stem}.json"
    
    if full_transcription_file.exists():
        log_info(f"Auto-detected full transcription for {task_id}")
        with open(full_transcription_file, 'r', encoding='utf-8') as f:
            cached_data = json.load(f)
            result = cached_data.get("result", [])
            status = "completed"
            progress = 100
    elif cache_file.exists():
        log_info(f"Auto-detected diarization for {task_id}")
        with open(cache_file, 'r') as f:
            timeline = json.load(f)
            status = "diarization_complete"
            progress = 100

    transcriptions[task_id] = {
        "filename": file.filename,
        "status": status,
        "progress": progress,
        "timeline": timeline,
        "result": result,
        "eta": "Waiting..."
    }
    
    return {"task_id": task_id, "status": status}

@app.post("/import_diarization/{task_id}")
async def import_diarization(task_id: str, timeline: list = Body(...)):
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    stem = Path(transcriptions[task_id]["filename"]).stem
    cache_file = CACHE_DIR / f"{stem}_diarize.json"
    with open(cache_file, 'w') as f:
        json.dump(timeline, f)
        
    transcriptions[task_id].update({
        "status": "diarization_complete",
        "progress": 100,
        "timeline": timeline
    })
    log_info(f"Manual diarization imported for {task_id}")
    return {"status": "success"}

@app.post("/import_transcription/{task_id}")
async def import_transcription(task_id: str, result: list = Body(...)):
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    transcriptions[task_id].update({
        "status": "completed",
        "progress": 100,
        "result": result
    })
    
    # Save to JSON, MD, DOCX
    stem = Path(transcriptions[task_id]["filename"]).stem
    full_json_path = UPLOAD_DIR / f"{stem}.json"
    with open(full_json_path, 'w', encoding='utf-8') as f:
        json.dump({"result": result}, f)
        
    regenerate_files(task_id)
    log_info(f"Manual transcription imported for {task_id}")
    return {"status": "success"}

@app.post("/diarize/{task_id}")
async def start_diarization(task_id: str):
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    file_path = UPLOAD_DIR / transcriptions[task_id]["filename"]
    t = threading.Thread(target=run_diarize_task, args=(file_path, task_id), daemon=True)
    t.start()
    return {"status": "started"}

@app.post("/transcribe/{task_id}")
async def start_transcription(task_id: str):
    if task_id not in transcriptions:
        return {"error": "Task not found"}
    
    file_path = UPLOAD_DIR / transcriptions[task_id]["filename"]
    t = threading.Thread(target=run_transcribe_task, args=(file_path, task_id), daemon=True)
    t.start()
    return {"status": "started"}


@app.get("/status/{task_id}")
async def get_status(task_id: str):
    return transcriptions.get(task_id, {"status": "not_found"})

@app.get("/audio/{filename}")
async def get_audio(filename: str):
    from fastapi.responses import FileResponse
    return FileResponse(UPLOAD_DIR / filename)

@app.get("/download/{filename}")
async def download_file(filename: str):
    from fastapi.responses import FileResponse
    path = UPLOAD_DIR / filename
    if path.exists():
        media_type = "text/markdown"
        if filename.endswith(".docx"):
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return FileResponse(path, media_type=media_type, filename=filename)
    return {"error": "File not found"}


class UpdateSpeakerRequest(BaseModel):
    task_id: str
    segment_index: int
    speaker_name: str

@app.post("/update_speaker")
async def update_speaker(req: UpdateSpeakerRequest):
    if req.task_id in transcriptions:
        task = transcriptions[req.task_id]
        if 0 <= req.segment_index < len(task["result"]):
            old_name = task["result"][req.segment_index]["speaker"]
            new_name = req.speaker_name
            
            # Bulk rename: update all segments that had this speaker
            for seg in task["result"]:
                if seg["speaker"] == old_name:
                    seg["speaker"] = new_name
                    
            # Regenerate files to reflect changes
            regenerate_files(req.task_id)
            return {"status": "success"}


    return {"status": "error", "message": "Task or segment not found"}

def regenerate_files(task_id):
    task = transcriptions[task_id]
    file_path = UPLOAD_DIR / task["filename"]
    
    # 1. Regenerate MD
    md_file_path = file_path.with_suffix(".md")
    md_content = f"# Transcription: {task['filename']}\n\n"
    for seg in task["result"]:
        md_content += f"**[{seg['timestamp']}] {seg['speaker']}:** {seg['text']}\n\n"
        
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    
    # 2. Regenerate DOCX
    generate_docx(task_id)


app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
