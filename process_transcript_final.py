
import docx
import re

def fix_text(text):
    # Dictionary of replacements for common transcription errors
    replacements = {
        r"комплоентность": "комплаентность",
        r"сахороснижающая": "сахароснижающая",
        r"полинеропатия": "полинейропатия",
        r"парастезии": "парестезии",
        r" УМС": " ОМС",
        r"Кринологи": "Эндокринологи",
        r"цифродавление": "цифры давления",
        r"симовик": "Семавик",
        r"земпик": "Оземпик",
        r"ситоглептин": "ситаглиптин",
        r"бисопролол а": "бисопролол А",
        r"не бевололом": "небивололом",
        r"гелок": "Эгилок",
        r"от НЛО": "Атенолол",
        r"в абрадин": "ивабрадин",
        r"бесопролол": "бисопролол",
        r"Престелол": "Престилол",
        r"ЭБС": "ИБС",
        r"ОГЭ": "АГ",
        r"ХСОН": "ХСН",
        r"битоблокатор": "бета-блокатор",
        r"Нишфарм": "Нижфарм",
        r"Кырка": "Крка",
        r"цифр давления": "цифры давления",
        r"цифра давления": "цифры давления",
        r"цифр сахара": "цифры сахара",
        r"по-всякому бывает": "Всегда бывает", # just an example of style adjust if needed, but I'll stick to errors
    }
    
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    
    # Capitalize the first letter of sentences if needed
    # (Usually whisper transcript is okay but sometimes it skips)
    
    return text

def process_transcript(input_path, output_path):
    doc = docx.Document(input_path)
    
    speaker_map = {
        "SPEAKER_01": "Интервьюер",
        "SPEAKER_00": "Респондент"
    }

    processed_blocks = []
    current_speaker = None
    current_text = ""
    current_time = ""

    pattern = re.compile(r"\[(\d{2}:\d{2}(?::\d{2})?)\] (SPEAKER_\d{2}): (.*)")

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        match = pattern.match(text)
        if match:
            time_str, speaker_id, content = match.groups()
            speaker_name = speaker_map.get(speaker_id, speaker_id)
            
            # Apply fix to content
            content = fix_text(content)
            
            if speaker_name == current_speaker:
                current_text += " " + content
            else:
                if current_speaker:
                    processed_blocks.append({
                        "speaker": current_speaker,
                        "time": current_time,
                        "text": current_text.strip()
                    })
                current_speaker = speaker_name
                current_time = time_str
                current_text = content
        else:
            if current_speaker:
                current_text += " " + fix_text(text)
            else:
                # Header lines
                processed_blocks.append({
                    "speaker": None,
                    "time": None,
                    "text": fix_text(text)
                })

    if current_speaker:
        processed_blocks.append({
            "speaker": current_speaker,
            "time": current_time,
            "text": current_text.strip()
        })

    new_doc = docx.Document()
    
    # Remove default spacing
    style = new_doc.styles['Normal']
    style.paragraph_format.space_after = 0
    style.paragraph_format.space_before = 0
    style.paragraph_format.line_spacing = 1.0

    for block in processed_blocks:
        p = new_doc.add_paragraph()
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0
        
        speaker = block["speaker"]
        time = block["time"]
        text = block["text"]
        
        if speaker == "Интервьюер":
            # Bold all words, No timeline
            run = p.add_run(f"Интервьюер: {text}")
            run.bold = True
        elif speaker == "Респондент":
            # Only "Респондент" bold, No timeline
            run_tag = p.add_run("Респондент")
            run_tag.bold = True
            run_rest = p.add_run(f": {text}")
        else:
            p.add_run(text)
            
    new_doc.save(output_path)
    return processed_blocks

if __name__ == "__main__":
    input_file = r"c:\Users\halfo\OneDrive\Documents\GitHub\whisperX-medical-transcribe\uploads\ГИ4_Тер_гос_Спб_КораблеваНВ_20.02.docx"
    output_file = r"c:\Users\halfo\OneDrive\Documents\GitHub\whisperX-medical-transcribe\uploads\ГИ4_Тер_гос_Спб_КораблеваНВ_20.02_fixed.docx"
    
    process_transcript(input_file, output_file)
    print("Processing complete.")
